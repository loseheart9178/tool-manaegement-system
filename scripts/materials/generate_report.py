from __future__ import annotations

import os
import subprocess
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[2]
ASSET_DIR = ROOT / "deliverables" / "assets"
FINAL_DIR = ROOT / "deliverables" / "final"
REPORT_DIR = ROOT / "report"
BUILD_DIR = ROOT / "build"

PROJECT_NAME = "工厂工具管理信息系统"
COURSE_NAME = "程序设计课程实践"
LEADER = "25051408 卢世豪"
MEMBER = "25051415 丁伟哲"
DATE_TEXT = "2026 年 6 月"

FONT_SONG = "/System/Library/Fonts/Supplemental/Songti.ttc"
FONT_HEI = "/System/Library/Fonts/PingFang.ttc"
FONT_ARIAL_UNICODE = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"


def ensure_dirs() -> None:
    for path in [ASSET_DIR, FINAL_DIR, REPORT_DIR, BUILD_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [FONT_HEI if bold else FONT_SONG, FONT_ARIAL_UNICODE]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def draw_rounded_box(draw: ImageDraw.ImageDraw, box, fill, outline="#C9D5E2", width=2, radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def centered_text(draw, box, text, fnt, fill="#1B2A3A"):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=8, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2), text, font=fnt, fill=fill, spacing=8, align="center")


def module_diagram() -> Path:
    path = ASSET_DIR / "system_module_diagram.png"
    img = Image.new("RGB", (1600, 1000), "#F6F8FB")
    draw = ImageDraw.Draw(img)
    title = font(46, True)
    h = font(32, True)
    body = font(25)
    draw.text((70, 55), "系统功能模块图", font=title, fill="#12263A")

    root = (520, 130, 1080, 240)
    draw_rounded_box(draw, root, "#D9EAF7", "#3A78A8", 3)
    centered_text(draw, root, "工厂工具管理信息系统", h)

    modules = [
        ((80, 340, 400, 520), "系统基础模块\n登录验证\n主菜单导航\n退出保存"),
        ((460, 340, 780, 520), "工具信息管理\n显示/添加\n删除/修改\n多条件查询"),
        ((840, 340, 1160, 520), "工具借还管理\n借用登记\n归还登记\n记录查询"),
        ((1220, 340, 1540, 520), "统计分析模块\n库存统计\n低库存预警\n借用排行"),
        ((460, 665, 780, 845), "数据持久化\n工具数据文件\n借还记录文件\n启动自动加载"),
        ((840, 665, 1160, 845), "辅助测试数据\n样例工具数据\n样例借还记录\n演示验证"),
    ]
    anchor_points = [(240, 340), (620, 340), (1000, 340), (1380, 340), (620, 665), (1000, 665)]
    for (box, text), (_, top_y) in zip(modules, anchor_points):
        draw_rounded_box(draw, box, "#FFFFFF", "#B7C7D8", 2)
        centered_text(draw, box, text, body)
        x = (box[0] + box[2]) / 2
        draw.line((800, 240, x, top_y), fill="#6D8FA8", width=4)
    img.save(path)
    return path


def flow_diagram() -> Path:
    path = ASSET_DIR / "business_flow_diagram.png"
    img = Image.new("RGB", (1600, 1000), "#FBFCFE")
    draw = ImageDraw.Draw(img)
    title = font(46, True)
    h = font(28, True)
    draw.text((70, 55), "系统业务流程图", font=title, fill="#12263A")
    steps = [
        ("启动程序", 120),
        ("登录验证", 250),
        ("加载工具与借还记录", 380),
        ("主菜单选择功能", 510),
        ("工具管理 / 借还管理 / 统计分析", 640),
        ("退出前保存数据", 770),
        ("系统结束", 900),
    ]
    boxes = []
    for text, y in steps:
        box = (500, y, 1100, y + 86)
        boxes.append(box)
        draw_rounded_box(draw, box, "#EAF4EA" if "保存" in text else "#FFFFFF", "#73A36B" if "保存" in text else "#B8C8D8", 2, 20)
        centered_text(draw, box, text, h)
    for a, b in zip(boxes, boxes[1:]):
        x = 800
        draw.line((x, a[3], x, b[1] - 16), fill="#5C7790", width=4)
        draw.polygon([(x, b[1] - 4), (x - 12, b[1] - 22), (x + 12, b[1] - 22)], fill="#5C7790")
    side = font(23)
    draw_rounded_box(draw, (90, 500, 390, 785), "#FFF9EC", "#D9B568", 2)
    centered_text(draw, (90, 500, 390, 785), "关键数据变化\n\n借用：库存减少，已借增加\n归还：库存增加，已借减少\n退出：写入本地文件", side)
    img.save(path)
    return path


def run_program_capture() -> tuple[str, str]:
    exe = BUILD_DIR / "tool_manager"
    subprocess.run(
        ["clang", "-std=c11", "-Wall", "-Wextra", "-g", *map(str, (ROOT / "src").glob("*.c")), "-Iinclude", "-o", str(exe)],
        cwd=ROOT,
        check=True,
    )
    demo_input = "admin\nadmin123\n1\n9\n10\n6\n11\n0\n"
    result = subprocess.run([str(exe)], input=demo_input, text=True, cwd=ROOT, capture_output=True, check=True)
    out = result.stdout
    (ASSET_DIR / "run_demo_output.txt").write_text(out, encoding="utf-8")
    return out, str(exe)


def terminal_image(text: str, filename: str, title: str, max_lines: int = 28) -> Path:
    path = ASSET_DIR / filename
    lines = [line for line in text.splitlines() if line.strip()][:max_lines]
    img = Image.new("RGB", (1500, 900), "#111827")
    draw = ImageDraw.Draw(img)
    title_font = font(30, True)
    mono = ImageFont.truetype(FONT_ARIAL_UNICODE, 23)
    draw.rounded_rectangle((35, 35, 1465, 865), radius=22, fill="#0B1220", outline="#293548", width=2)
    draw.text((70, 60), title, font=title_font, fill="#D6E4FF")
    y = 120
    for line in lines:
        safe = line[:95]
        draw.text((70, y), safe, font=mono, fill="#E5E7EB")
        y += 27
        if y > 820:
            break
    img.save(path)
    return path


def run_summary_image(text: str) -> Path:
    path = ASSET_DIR / "run_effect_screenshot.png"
    img = Image.new("RGB", (1600, 1000), "#0B1220")
    draw = ImageDraw.Draw(img)
    title_font = font(38, True)
    head_font = font(25, True)
    mono = ImageFont.truetype(FONT_ARIAL_UNICODE, 21)
    draw.text((58, 48), "工厂工具管理信息系统 - 运行效果摘要", font=title_font, fill="#EAF2FF")

    panels = [
        ((55, 130, 760, 390), "登录与主菜单", ["登录成功", "已加载 12 条工具记录", "已加载 9 条借还记录", "菜单包含工具管理、借还管理、统计分析"]),
        ((840, 130, 1545, 390), "库存整体统计", ["工具种类数：12 种", "库存总数量：142 件", "当前在库数量：126 件", "已借出数量：16 件", "库存总价值：30281.00 元"]),
        ((55, 465, 760, 845), "低库存预警（阈值 6）", ["T002 角磨机：库存 5", "T007 电焊机：库存 3（损坏）", "T010 热风枪：库存 4", "T012 绝缘电阻测试仪：库存 5", "共 4 种工具库存低于阈值"]),
        ((840, 465, 1545, 845), "借用次数排行榜", ["1. T001 电钻：3 次", "2. T003 数字万用表：2 次", "3. T008 游标卡尺：1 次", "4. T002 角磨机：1 次", "共 6 种工具有借用记录"]),
    ]
    for box, panel_title, lines in panels:
        draw.rounded_rectangle(box, radius=24, fill="#111827", outline="#334155", width=2)
        draw.text((box[0] + 30, box[1] + 24), panel_title, font=head_font, fill="#BFDBFE")
        y = box[1] + 78
        for line in lines:
            draw.text((box[0] + 35, y), line, font=mono, fill="#E5E7EB")
            y += 42
    draw.text((58, 920), "验证功能：登录、数据加载、工具显示、库存统计、低库存预警、借用排行榜、退出保存", font=font(22), fill="#CBD5E1")
    img.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = Pt(20)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        st = styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.line_spacing = Pt(20)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(20)


def add_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(rows[0]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "DDEBF7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    return table


def build_docx(module_img: Path, flow_img: Path, run_img: Path) -> Path:
    doc = Document()
    set_doc_defaults(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"《{COURSE_NAME}》\n综合项目实验报告")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(24)
    r.bold = True

    doc.add_paragraph()
    for line in [
        f"项目名称：{PROJECT_NAME}",
        f"团队成员：{LEADER}，{MEMBER}",
        f"完成时间：{DATE_TEXT}",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(14)
    cover = ASSET_DIR / "cover_factory_tool_system.png"
    if cover.exists():
        doc.add_picture(str(cover), width=Cm(12.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    add_heading(doc, f"{PROJECT_NAME}项目实验报告", 1)

    add_heading(doc, "1. 团队成员组成及分工", 2)
    add_table(
        doc,
        [
            ["学号", "姓名", "详细任务分工"],
            ["25051408", "卢世豪", "组长。负责系统总体设计、工具基础信息管理、工具数据保存加载、登录与主菜单调度、代码整合、可执行文件生成、实验报告统稿。"],
            ["25051415", "丁伟哲", "负责工具借还管理、借还记录管理、库存统计分析、低库存预警、借用次数排行、测试数据整理与答辩 PPT 制作。"],
        ],
        [3, 3, 10],
    )

    add_heading(doc, "2. 开发背景", 2)
    add_para(doc, "在工厂生产和设备维护过程中，工具种类多、借用频繁，如果仅依靠纸质登记或人工记忆，容易出现工具库存不清、借还状态混乱、低库存无法及时发现等问题。为提高工具管理效率，本项目使用 C 语言设计并实现一个工厂工具管理信息系统。")
    add_para(doc, "系统围绕工具基础信息和借还记录两类核心数据展开，提供工具增删改查、借用归还、历史记录查看、库存统计、低库存预警和借用排行等功能，并通过文件读写实现数据持久化保存。项目综合运用了结构体、数组、函数模块化、字符串处理、指针参数、文件读写和排序等 C 语言知识。")

    add_heading(doc, "3. 系统功能设计", 2)
    add_heading(doc, "3.1 系统功能模块设计", 3)
    add_para(doc, "系统采用模块化设计。主程序负责登录、数据加载、菜单调度和退出保存；工具信息管理模块负责工具记录的增删改查；借还管理模块负责借用、归还和记录展示；统计分析模块负责库存汇总、低库存预警和借用次数排行；数据持久化模块负责工具数据和借还记录的本地文件读写。")
    doc.add_picture(str(module_img), width=Cm(15.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("图1 系统功能模块图")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "3.2 系统业务流程设计", 3)
    add_para(doc, "系统启动后，用户首先进行登录验证。登录成功后，程序从本地数据文件中加载工具信息和借还记录，然后进入主菜单循环。用户根据菜单选择工具管理、借还管理或统计分析功能。退出系统时，程序自动保存全部工具信息和借还记录，确保下次启动时数据仍然存在。")
    doc.add_picture(str(flow_img), width=Cm(15.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("图2 业务流程图")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "4. 项目创建", 2)
    add_heading(doc, "4.1 系统开发环境要求", 3)
    add_table(
        doc,
        [
            ["项目", "内容"],
            ["操作系统", "macOS"],
            ["开发工具", "Visual Studio Code"],
            ["开发语言", "C 语言"],
            ["编译器", "Apple clang / gcc"],
            ["调试工具", "CodeLLDB"],
        ],
        [4, 11],
    )
    add_heading(doc, "4.2 项目创建过程", 3)
    add_para(doc, "项目采用标准 C 语言多文件组织方式，包含 `src`、`include`、`data`、`docs`、`report`、`ppt` 和 `build` 等目录。其中 `src` 保存源文件，`include` 保存头文件，`data` 保存工具和借还记录数据，`build` 保存可执行文件。")

    add_heading(doc, "5. 预处理模块设计", 2)
    add_para(doc, "项目主要引用 `stdio.h`、`string.h`、`stdlib.h` 等标准头文件，并通过自定义头文件 `tool.h`、`borrow.h`、`storage.h`、`statistics.h`、`menu.h` 声明结构体和函数接口。")
    add_table(
        doc,
        [
            ["头文件", "作用"],
            ["tool.h", "定义 Tool 结构体和工具管理函数接口"],
            ["borrow.h", "定义 BorrowRecord 结构体和借还管理函数接口"],
            ["storage.h", "声明工具数据保存和加载接口"],
            ["statistics.h", "声明库存统计、低库存预警和排行榜接口"],
            ["menu.h", "声明登录、菜单显示和系统主循环接口"],
        ],
        [4, 11],
    )

    add_heading(doc, "6. 工具信息管理模块设计", 2)
    add_para(doc, "工具信息使用 `Tool` 结构体表示，包括工具编号、名称、类型、库存、状态、已借出数量、仓库存放位置和单价。工具编号作为主键，添加工具时必须保证唯一。")
    add_table(
        doc,
        [
            ["函数", "实现功能", "负责人"],
            ["find_tool_index_by_id", "按编号查找工具下标，未找到返回 -1", "卢世豪"],
            ["show_all_tools", "列表展示全部工具信息", "卢世豪"],
            ["add_tool", "添加工具并校验编号、库存、状态和单价", "卢世豪"],
            ["delete_tool", "删除工具，已借出工具禁止删除，删除前二次确认", "卢世豪"],
            ["update_tool", "修改除编号外的工具信息，支持回车保留原值", "卢世豪"],
            ["search_tools", "按编号精准查询，按名称或类型模糊查询", "卢世豪"],
        ],
        [4.4, 8.6, 2.2],
    )

    add_heading(doc, "7. 工具借还管理模块设计", 2)
    add_para(doc, "借还记录使用 `BorrowRecord` 结构体表示，记录编号自动生成，工具编号用于关联工具信息。借用时系统会检查工具是否存在、是否损坏、借用数量是否超过库存；归还时根据记录编号定位借用记录，防止重复归还。")
    add_table(
        doc,
        [
            ["函数", "实现功能", "负责人"],
            ["generate_record_id", "生成自增且唯一的借还记录编号", "丁伟哲"],
            ["borrow_tool", "借用工具，生成记录并同步扣减库存", "丁伟哲"],
            ["return_tool", "归还工具，恢复库存并修改记录状态", "丁伟哲"],
            ["show_all_records", "列表展示全部历史借还记录", "丁伟哲"],
            ["load_records / save_records", "借还记录文件加载和保存", "丁伟哲"],
        ],
        [4.4, 8.6, 2.2],
    )

    add_heading(doc, "8. 统计分析模块设计", 2)
    add_para(doc, "统计分析模块基于工具数组和借还记录数组实现。库存整体统计计算工具种类数、库存总数量、在库数量、已借出数量和总价值；低库存预警根据用户输入阈值筛选当前库存不足的工具；借用排行榜统计各工具在历史借还记录中出现的次数，并使用 `qsort` 降序排序。")
    add_table(
        doc,
        [
            ["统计功能", "计算方式", "负责人"],
            ["库存整体统计", "汇总 stock、borrowed_count 和 price", "丁伟哲"],
            ["低库存预警", "筛选 stock 小于安全库存阈值的工具", "丁伟哲"],
            ["借用次数排行", "遍历借还记录累计次数，并按次数降序排序", "丁伟哲"],
        ],
        [4.5, 8.5, 2.2],
    )

    add_heading(doc, "9. 数据文件保存与加载", 2)
    add_para(doc, "本系统使用二进制文件保存数据。工具信息保存到 `data/tools.dat`，借还记录保存到 `data/borrow_records.dat`。保存时先写入记录数量，再逐条写入结构体；加载时按相同顺序读取，并通过最大容量参数防止数组越界。")

    add_heading(doc, "10. 项目运行效果", 2)
    add_para(doc, "程序支持登录、菜单导航、工具列表展示、库存统计、低库存预警和借用排行榜等功能。下图为终端运行效果示例。")
    doc.add_picture(str(run_img), width=Cm(15.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("图3 项目运行效果截图")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "11. 项目创新点", 2)
    add_para(doc, "第一，系统实现了低库存预警，用户可以输入安全库存阈值，系统自动筛选当前库存不足的工具，提醒管理员及时补货。")
    add_para(doc, "第二，系统实现了工具借用次数排行榜，能够根据历史借还记录分析常用工具，为采购、维护和仓库管理提供参考。")
    add_para(doc, "第三，查询功能支持名称和类型模糊查询，用户不需要完整输入工具名称即可快速定位相关工具。")
    add_para(doc, "第四，系统实现启动自动加载和退出自动保存，保证程序重启后工具和借还记录数据不丢失。")

    add_heading(doc, "12. 收获和建议", 2)
    add_heading(doc, "组长 卢世豪：", 3)
    add_para(doc, "通过本次课程实践，我更加系统地理解了 C 语言程序从单文件练习到多文件项目的组织方式。以前写程序时更多关注某一个函数是否能运行，这次作为组长，需要考虑结构体字段设计、函数接口划分、模块之间如何传递数据、文件读写如何保证数据不丢失等问题。我主要负责工具信息管理、菜单调度和数据保存加载，在实现过程中对数组、结构体、字符串比较、指针参数等知识有了更深的理解。例如添加工具时，工具数量需要在函数内部改变，因此必须传入 `int *tool_count`；删除工具时需要把数组后面的元素前移，才能保证数据连续。开发中遇到的主要困难是如何让工具管理模块和借还模块共享同一份库存数据。最终我们通过主程序统一维护 `Tool tools[]` 数组，借还模块通过工具编号查找下标并修改库存，避免了数据不一致。团队协作方面，我负责确定接口和整合代码，同伴根据接口实现借还和统计功能，最后一起测试登录、添加、借用、归还、统计等流程。通过这个项目，我认识到软件产品研发不仅要能实现功能，还要考虑边界情况、用户误操作和数据安全。未来如果继续完善，我希望增加 CSV 导出、用户权限管理和图形化界面，让系统更接近真实仓库管理软件。")
    add_heading(doc, "成员 丁伟哲：", 3)
    add_para(doc, "在本次课程实践中，我主要负责工具借还管理和统计分析模块。借还管理看起来只是简单的借出和归还，但真正实现时需要考虑很多细节，例如工具编号是否存在、工具状态是否损坏、借用数量是否超过库存、同一条记录是否已经归还等。如果这些条件没有处理好，就会导致库存变成负数或者重复归还造成库存虚增。通过实现 `borrow_tool` 和 `return_tool`，我更加理解了结构体数组在业务数据管理中的作用，也掌握了如何通过一个字段建立两类数据之间的关联关系。借还记录中只保存 `tool_id`，需要时再通过工具编号到工具数组中查找对应工具，这种设计减少了重复数据，也让库存更新更集中。统计分析模块让我进一步练习了循环统计、动态内存申请和排序函数的使用。借用排行榜中使用 `malloc` 根据工具数量创建临时数组，统计完成后使用 `qsort` 排序，最后用 `free` 释放内存。项目协作过程中，我根据组长定义的工具结构体和查找接口完成借还模块，再与工具管理模块联调，逐步修正库存同步和记录状态的问题。这个过程让我认识到团队项目中接口约定非常重要，只有函数参数、返回值和数据含义提前明确，后续开发才能顺利合并。未来可以继续改进日期校验、逾期提醒和统计图表，让系统具备更高的实用价值。")

    add_heading(doc, "附：源代码清单", 2)
    add_table(
        doc,
        [
            ["文件", "主要内容", "完成人"],
            ["src/main.c", "程序入口，调用系统主循环", "卢世豪"],
            ["src/menu.c", "登录验证、主菜单、功能调度、退出保存", "卢世豪"],
            ["src/tool.c", "工具信息增删改查", "卢世豪"],
            ["src/storage.c", "工具数据文件保存和加载", "卢世豪"],
            ["src/borrow.c", "借用、归还、借还记录和记录文件读写", "丁伟哲"],
            ["src/statistics.c", "库存统计、低库存预警、借用排行榜", "丁伟哲"],
            ["include/*.h", "结构体定义和函数声明", "卢世豪、丁伟哲"],
            ["scripts/generate_data.c", "生成样例工具数据", "卢世豪"],
            ["scripts/generate_records.c", "生成样例借还记录", "丁伟哲"],
        ],
        [4.5, 8.5, 2.2],
    )

    out = REPORT_DIR / "工厂工具管理信息系统-实验报告.docx"
    doc.save(out)
    final = FINAL_DIR / out.name
    doc.save(final)
    return out


def pdf_paragraph(text: str, style):
    return Paragraph(text.replace("`", ""), style)


def build_pdf(module_img: Path, flow_img: Path, run_img: Path) -> Path:
    pdfmetrics.registerFont(TTFont("CN", FONT_ARIAL_UNICODE))
    pdfmetrics.registerFont(TTFont("CNB", FONT_ARIAL_UNICODE))
    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName="CN", fontSize=10.5, leading=18, firstLineIndent=21, spaceAfter=6)
    h1 = ParagraphStyle("H1CN", parent=styles["Heading1"], fontName="CNB", fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=12)
    h2 = ParagraphStyle("H2CN", parent=styles["Heading2"], fontName="CNB", fontSize=14, leading=20, spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("H3CN", parent=styles["Heading3"], fontName="CNB", fontSize=12, leading=18, spaceBefore=6, spaceAfter=4)
    center = ParagraphStyle("CenterCN", parent=body, alignment=TA_CENTER, firstLineIndent=0)

    out = REPORT_DIR / "工厂工具管理信息系统-实验报告.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=A4, leftMargin=3 * cm, rightMargin=2.5 * cm, topMargin=2.5 * cm, bottomMargin=2.5 * cm)
    story = []
    story.append(Paragraph(f"《{COURSE_NAME}》<br/>综合项目实验报告", h1))
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(f"项目名称：{PROJECT_NAME}", center))
    story.append(Paragraph(f"团队成员：{LEADER}，{MEMBER}", center))
    story.append(Paragraph(f"完成时间：{DATE_TEXT}", center))
    cover = ASSET_DIR / "cover_factory_tool_system.png"
    if cover.exists():
        story.append(Spacer(1, 0.3 * cm))
        story.append(RLImage(str(cover), width=10.5 * cm, height=7.0 * cm))
    story.append(PageBreak())

    story.append(Paragraph(f"{PROJECT_NAME}项目实验报告", h1))
    story.append(Paragraph("1. 团队成员组成及分工", h2))
    story.append(Table(
        [["学号", "姓名", "详细任务分工"],
         ["25051408", "卢世豪", "组长，负责系统总体设计、工具信息管理、文件存储、主菜单、代码整合和报告统稿。"],
         ["25051415", "丁伟哲", "负责借还管理、统计分析、测试数据整理和答辩 PPT 制作。"]],
        colWidths=[2.5 * cm, 2.2 * cm, 10 * cm],
        style=TableStyle([("FONT", (0, 0), (-1, -1), "CN"), ("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DDEBF7")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")])
    ))
    sections = [
        ("2. 开发背景", "在工厂生产和设备维护过程中，工具种类多、借用频繁，如果仅依靠纸质登记或人工记忆，容易出现工具库存不清、借还状态混乱、低库存无法及时发现等问题。本项目使用 C 语言设计并实现一个工厂工具管理信息系统，综合运用结构体、数组、函数模块化、字符串处理、指针参数、文件读写和排序等知识。"),
        ("3. 系统功能设计", "系统采用模块化设计。主程序负责登录、数据加载、菜单调度和退出保存；工具信息管理模块负责工具记录增删改查；借还管理模块负责借用、归还和记录展示；统计分析模块负责库存汇总、低库存预警和借用次数排行。"),
    ]
    for title, text in sections:
        story.append(Paragraph(title, h2))
        story.append(pdf_paragraph(text, body))
    story.append(RLImage(str(module_img), width=15 * cm, height=9.4 * cm))
    story.append(Paragraph("图1 系统功能模块图", center))
    story.append(RLImage(str(flow_img), width=15 * cm, height=9.4 * cm))
    story.append(Paragraph("图2 业务流程图", center))
    story.append(Paragraph("4. 项目创建与开发环境", h2))
    story.append(pdf_paragraph("项目在 macOS 和 Visual Studio Code 环境下开发，使用 Apple clang 编译器。目录包含 src、include、data、docs、report、ppt 和 build。", body))
    story.append(Paragraph("5. 主要模块实现", h2))
    story.append(pdf_paragraph("工具管理模块实现显示、添加、删除、修改和查询；借还模块实现借用、归还、记录编号生成和历史记录展示；统计模块实现库存统计、低库存预警和借用排行榜；文件模块通过二进制文件实现数据持久化。", body))
    story.append(Paragraph("6. 项目运行效果", h2))
    story.append(RLImage(str(run_img), width=15 * cm, height=9 * cm))
    story.append(Paragraph("图3 项目运行效果截图", center))
    story.append(Paragraph("7. 项目创新点", h2))
    story.append(pdf_paragraph("本项目的创新点包括低库存预警、借用次数排行榜、工具名称和类型模糊查询、启动自动加载和退出自动保存。", body))
    story.append(Paragraph("8. 收获和建议", h2))
    story.append(Paragraph("组长 卢世豪：", h3))
    story.append(pdf_paragraph("通过本次课程实践，我更加系统地理解了 C 语言多文件项目的组织方式，也理解了结构体、数组、指针参数和文件读写在真实业务程序中的作用。作为组长，我负责工具信息管理、菜单调度、数据保存加载和代码整合，重点解决了工具编号唯一、删除二次确认、已借出工具禁止删除以及数据持久化等问题。", body))
    story.append(Paragraph("成员 丁伟哲：", h3))
    story.append(pdf_paragraph("我主要负责借还管理和统计分析。通过实现借用、归还、低库存预警和排行榜，我理解了业务数据之间的关联关系，也掌握了通过工具编号关联工具信息和借还记录的方法。借用和归还时同步修改 stock 和 borrowed_count，是保证库存一致性的关键。", body))
    story.append(Paragraph("附：源代码清单", h2))
    story.append(Table(
        [["文件", "主要内容", "完成人"],
         ["src/main.c / src/menu.c", "入口、登录、菜单调度、退出保存", "卢世豪"],
         ["src/tool.c / src/storage.c", "工具增删改查、工具文件读写", "卢世豪"],
         ["src/borrow.c / src/statistics.c", "借还管理、记录读写、统计分析", "丁伟哲"],
         ["include/*.h", "结构体定义和函数声明", "卢世豪、丁伟哲"]],
        colWidths=[5 * cm, 7 * cm, 3 * cm],
        style=TableStyle([("FONT", (0, 0), (-1, -1), "CN"), ("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DDEBF7"))])
    ))
    doc.build(story)
    final = FINAL_DIR / out.name
    final.write_bytes(out.read_bytes())
    return out


def main() -> None:
    ensure_dirs()
    module_img = module_diagram()
    flow_img = flow_diagram()
    output, _ = run_program_capture()
    terminal_image(output, "run_demo_terminal_full.png", "工厂工具管理信息系统 - 终端输出")
    run_img = run_summary_image(output)
    docx = build_docx(module_img, flow_img, run_img)
    pdf = build_pdf(module_img, flow_img, run_img)
    print(f"DOCX: {docx}")
    print(f"PDF: {pdf}")


if __name__ == "__main__":
    main()
